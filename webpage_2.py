from flask import Flask, request, jsonify, render_template, send_file, session, redirect, url_for
from fpdf import FPDF
import os
import uuid
import json
from chatbot import generate_openai_response,generate_steps,generate_openai_summary,generate_state,generate_openai_response_tenured,generate_plot_code,document_chat,extract_kyc_info_v2
import pandas as pd
from docx import Document
import docx
import sys
import os
from app_v2 import app_v2,extract_summary
from least_impact import least_impact
import pandas as pd
from question_formatted import format_question
from chart_prompt import plot_type_prompt_generation
from chart_prompt import plot_fixed_chart
from langchain_openai import ChatOpenAI
import httpx
import plotly.express
import plotly.express as px
import plotly.graph_objects as go

# Add the root directory to sys.path (assuming you're running the script from folder2)
app = Flask(__name__)

# Directory to save uploaded images
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# In-memory storage for conversation history
conversation = []
user_details = {}
CSV_FILE_PATH = 'incident_details.csv'

app.secret_key = 'asdfghjkl'  # Replace with a strong secret key

# Simulated user database (replace with a proper database in production)
users = {
    "Naman": "naman",
    "Alex": "alex"
}

@app.route('/')
def home():
    if 'username' in session:
        username = session['username']

        # Check if username is 'Naman' to trigger the prompt
        #if username == 'Naman':
            #return render_template('index_naman.html', username=session['username'])
        #else:
        return render_template('index.html', username=session['username'])
    return redirect(url_for('login'))

@app.route('/login', methods=['GET'])
def login():
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login_post():
    global current_username
    data = request.json
    username = data.get('username')
    password = data.get('password')
    current_username = username

    if username in users and users[username] == password:
        session['username'] = username
        return jsonify({'message': 'Login successful!'}), 200
    return jsonify({'message': 'Invalid username or password.'}), 401

@app.route('/logout', methods=['POST'])
def logout():
    session.pop('username', None)
    return jsonify({'message': 'Logged out successfully!'}), 200

@app.route('/get_alert_types', methods=['GET'])
def get_alert_types():
    df = pd.read_csv(CSV_FILE_PATH)
    incident_numbers = df['incident_number'].tolist()  # Adjust column name as necessary
    #print(incident_numbers)
    incident_numbers_v2 = ['INC196188','INC196104','INC196099','INC195996']
    return jsonify(incident_numbers_v2)

@app.route('/generate_report', methods=['POST'])
def generate_report():
    data = request.json
    report_input = data.get('report')
    try:
        # Here, add the logic to process the report_input
        # For example, call the OpenAI API to generate a summary or report
        response = f"Generated report based on: {report_input}"  # Replace with actual processing logic
        return jsonify({'response': response})
    except Exception as e:
        print(f"Error generating report: {e}")
        return jsonify({'error': 'Failed to generate report.'}), 500

from app_v2 import read_csv_errors
current_incident_number = None
@app.route('/load_data', methods=['GET'])
def load_data():
    alert_type = request.args.get('alert_type')
    selected_table = request.args.get('table')
    global current_incident_number
    df_2 = pd.read_csv('database/output.csv')
    df_3 = df_2[df_2['incident_number'] == alert_type]
    if alert_type and current_incident_number != alert_type:
        current_incident_number = alert_type
        read_csv_errors(df_3)
    # Map tables to CSV file paths
    csv_file_map = {
        'Incident Data': 'incident_details.csv',
        'Transaction Report': 'Transaction Report Sample.csv',
        'Attack Taxonomy':'Attack_taxonomy.csv',
        'Declines' : 'Declines.csv',
        'Impact' : 'impact.csv',
        'P History': 'P_history.csv'
    }

    # Load the appropriate CSV based on the selected table
    csv_file_path = csv_file_map.get(selected_table)
    def parse_json(json_str):
        return json.loads(json_str)
    if not csv_file_path:
            return jsonify({'error': 'Invalid table selected'}), 400
    if csv_file_path == 'incident_details.csv':
    # Load the CSV file into a DataFrame
            df = pd.read_csv(csv_file_path)
    # Filter DataFrame based on the alert type (assuming 'incident_number' column exists)
            filtered_df = df[df['incident_number'] == alert_type]
            json_data = pd.json_normalize(filtered_df['json_data'].apply(parse_json))
            melted_data = json_data.melt(var_name='Metric', value_name='value')
    # Convert the DataFrame to HTML
            table = melted_data.to_html(classes='table table-striped', index=False)
            return jsonify({'table': table})
    elif csv_file_path == 'Transaction Report Sample.csv':
        df = pd.read_csv(csv_file_path)
        filtered_df = df[df['incident number'] == alert_type]
        filtered_df = filtered_df.drop(columns=['incident number'])
        #melted_data = df.melt(var_name='variable', value_name='value')
        table = filtered_df.head(15).to_html(classes='table table-striped', index=False)
        return jsonify({'table': table})
    else:
        df = pd.read_csv(csv_file_path)
        filtered_df = df[df['incident number'] == alert_type]
        filtered_df = filtered_df.drop(columns=['incident number'])
        #melted_data = df.melt(var_name='variable', value_name='value')
        table = filtered_df.to_html(classes='table table-striped', index=False)
        return jsonify({'table': table})
    
# Route to get recommended questions
@app.route('/get_recommended_questions', methods=['GET'])
def get_recommended_questions():
    alert_type = request.args.get('alert_type')
    # Load questions from a CSV file
    df = pd.read_csv('questions.csv', encoding='latin1')  # Adjust the path as necessary
    filtered_df = df[df['incident number'] == alert_type]
    questions = filtered_df['Questions'].tolist()  # Assuming there's a column named 'Question'
    print(questions)
    return jsonify(questions)
# Existing route for asking queries
@app.route('/ask_query', methods=['POST'])
def ask_query():
    data = request.json
    query = data.get('query')
    alert_type = request.args.get('alert_type')
    df = pd.read_csv('questions.csv', encoding='latin1')
    df = df[df['incident number'] == alert_type]
    matching_row = df[df['Questions'] == query]
    if not matching_row.empty:
        answer = matching_row['Answer'].values[0]  # Get the answer
    elif matching_row.empty:
        ans, df, sql_query, chart_json = app_v2(query)
# Generate the summary (this function should extract summary from 'ans')
        summary = extract_summary(df)
        formatted_summary = summary.replace("\n", "<br>")

        
        
    
    # Generate an empty Plotly figure (no bars, no data)
        


        
        
        
# Initialize the HTML table with the SQL query
        
# Create the HTML table from the DataFrame using df.to_html()
# Add the formatted summary after the table
        table_html = '''
    <style>
        table {
            max-width: 100%;
            overflow-x: auto;
            word-wrap: break-word;
        }
        .content-wrapper {
            max-width: 100%;
            overflow-x: auto;
            word-wrap: break-word;
            padding: 10px;
        }
        .summary {
            max-width: 100%;
            overflow-x: auto;
            word-wrap: break-word;
            white-space: normal;  /* Ensures summary text wraps properly */
            padding: 10px;
        }
    </style>
'''
# Add the wrapper div for both table and summary
        table_html += '<div class="content-wrapper">'
# Add the table HTML
        table_html += df.to_html(classes='table table-striped', index=False)
# Add the summary after the table with custom class for better styling
        table_html += f'<p class="summary"><strong>Summary:</strong> {formatted_summary}</p>'
# Close the wrapper div
        table_html += '</div>'
        return jsonify({
                "answer": table_html,
                "chart_data": chart_json
            })
    else:
        answer = "Sorry, I don't have an answer for that question."
    return jsonify(answer=answer)

global_steps = None  # Initialize a global variable to store steps

@app.route('/get_flow_steps', methods=['GET'])
def get_flow_steps():
    global global_steps  # Access the global variable
    alert_type = request.args.get('alert_type')
    print(f'Alert Type: {alert_type}')  # Debug print
    steps = generate_steps(alert_type)
    
    if steps:
        global_steps = steps  # Store the steps globally
    else:
        print('No steps generated.')  # Debug print for empty steps

    sentences_df = pd.DataFrame(steps, columns=["Steps"])
    steps_table = sentences_df.to_html(classes='table table-striped', index=False)
    return jsonify({'table1': steps_table})

@app.route('/check_step_state', methods=['POST'])
def check_step_state():
    global global_steps  # Access the global steps variable
    if global_steps:
        # Get the completion state for the steps
        state = generate_state(global_steps)
    else:
        state = ['Please load flow steps']

    # Split the state entries into two columns
    split_states = []
    for entry in state:
        if ' - ' in entry:
            step, status = entry.split(' - ', 1)  # Split on the first occurrence of ' - '
            split_states.append({'Step': step.strip(), 'State': status.strip()})
        else:
            split_states.append({'Step': entry.strip(), 'State': 'Unknown'})  # Handle unexpected format

    # Create the HTML table manually with row coloring
    table_rows = []
    for row in split_states:
        if row['State'].lower() == 'completed':
            color = 'background-color: #d4edda;'  # Light Green
        elif row['State'].lower() == 'ongoing':
            color = 'background-color: #fff3cd;'  # Light Yellow
        elif row['State'].lower() == 'not completed':
            color = 'background-color: #f8d7da;'  # Light Red
        else:
            color = 'background-color: white;'  # Default color for unknown states

        table_rows.append(f'<tr style="{color}"><td>{row["Step"]}</td><td>{row["State"]}</td></tr>')

    # Combine rows into a complete HTML table
    state_df_html = f'''
    <table class="table table-striped">
        <thead>
            <tr>
                <th>Step</th>
                <th>State</th>
            </tr>
        </thead>
        <tbody>
            {''.join(table_rows)}
        </tbody>
    </table>
    '''

    return jsonify({'state': state_df_html})

@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    global summary
    summary_list = []

    df_attack = least_impact("attack taxonomy")
    max_decline_count = df_attack['Declined_Transaction_Count'].max()
    df = df_attack[df_attack['Declined_Transaction_Count'] == max_decline_count]
    df = df.head(1)
    attack_summary = f"1. Testing is ongoing on POS Entry Mode {df['POS_Entry_Mode'].values[0]}/ POS Condition Code {df['POS_Condition_Code'].values[0]}/ {df['Country_Current'].values[0]}/ ${df['Transaction_Amount_USD'].values[0]}.\n"
    summary_list.append(attack_summary)

    df_response = least_impact("calculate top response codes")
    df = df_response.sort_values(by = 'Occurrences', ascending = False).head(3)
    top_code = list(df['Authorization_Response_Code'].values)

    response_summary = f"2. Top Response Codes are {top_code[0]}, {top_code[1]} , {top_code[2]}.\n"
    summary_list.append(response_summary)

    mit_df = least_impact("Merchant Initiated transactions")
    mit_0 = mit_df[mit_df['Merchant_Initiated_Transaction_Classification'] == 0]
    mit_6 = mit_df[mit_df['Merchant_Initiated_Transaction_Classification'] == 6]

    mit_summary = f"3. For MIT code is equal to 0 we have {mit_0['count_of_declined_transactions'].values[0]} count of declined transactions and MIT code is equal to 6 we have {mit_6['count_of_declined_transactions'].values[0]} count of declined transactions.\n"
    summary_list.append(mit_summary)

    
    mcc_df = least_impact("different type of MCCs are available")
    mcc_code = list(mcc_df['Merchant_Category_Code'].values)
    mcc_value = list(mcc_df['Merchant_Code_Description'].values)

    mcc_summary = "4. Different type of MCCs available are "
    for i in range(len(mcc_code)):
        if i != len(mcc_code)-1:
            mcc_summary = mcc_summary + f"{mcc_code[i]}({mcc_value[i]})" + str(",") + str(" ")
        else:
            mcc_summary = mcc_summary + f"{mcc_code[i]}({mcc_value[i]})" + str(".\n")
    summary_list.append(mcc_summary)

    region_df = least_impact("domestic or international decline transactions count")

    domestic = region_df[region_df['Transaction_Type'] == 'Domestic']
    international = region_df[region_df['Transaction_Type'] == 'International']
    if domestic.empty:
        domestic_df = pd.DataFrame({'Transaction_Type': ['Domestic'], 'Declined_Transaction_Count': [0]})
        domestic = domestic_df[domestic_df['Transaction_Type'] == 'Domestic']

    if international.empty:
        international_df = pd.DataFrame({'Transaction_Type': ['International'], 'Declined_Transaction_Count': [0]})
        international = international_df[international_df['Transaction_Type'] == 'International']

    region_summary = f"5. We have {domestic['Declined_Transaction_Count'].values[0]} decline transaction counts for domeatic and {international['Declined_Transaction_Count'].values[0]} decline transaction counts for international.\n"
    summary_list.append(region_summary)

    holding_df = least_impact("multiple merchant name on same acquirer bin then it is called holding caid")
    if holding_df.empty:
        holding_summary  = f"6. There is not any holding caid available.\n"
    else:
        holding_summary  = f"6. There is holding caid available for acquirer bin({holding_df['Acquiring_Identifier_1'].values[0]}).\n"

    summary_list.append(holding_summary)

    print("Received request for summary")  # Check if the route is hit
    summary = ''
    for i in summary_list:
        summary = summary + i
    
    # Format the summary for HTML display
    summary_html = summary.replace('\n', '<br>').replace('•', '<li>').replace('\n', '</li><li>')
    summary_html = f'<ul><li>{summary_html}</li></ul>'
    
    return jsonify({'summary': summary_html})

@app.route('/chat', methods=['POST'])
def chat():
    global user_details
    
    user_name = request.form.get('user_name')
    user_id = request.form.get('employee_id')
    user_input = request.form.get('user_input')
    json_input = request.form.get('json_input')
    image = request.files.get('user_image')

    if user_name and user_id:
        user_details = {
            'user_name': user_name,
            'employee_id': user_id,
        }

    if json_input:
        try:
            data = json.loads(json_input)
            json_data = json.dumps(data)
        except json.JSONDecodeError:
            conversation.append("Invalid JSON format")
    
    conversation.append(f"User Input: {user_input}")
    #if current_username == 'Naman':
        #response, alert_type, json_data,incident_number = generate_openai_response_tenured(user_input)
    #elif current_username == 'Alex':
    response, alert_type, json_data,incident_number = generate_openai_response(user_input)
    if json_data:
        try:
            data = json.loads(json_data)
            json_data_v2 = json.dumps(data)
        except json.JSONDecodeError:
            conversation.append("Invalid JSON format")
    alert_type = alert_type + "Block"
    conversation.append(f"Alex: {response}")
    image_filename = None
    if image:
        image_filename = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}_{image.filename}")
        image.save(image_filename)
        conversation.append(f"Image: {image_filename}")
    answer_html = response.replace('\n', '<br>')
    print(incident_number)
    return jsonify({'response': answer_html,'alert_type':alert_type, 'json_data': json_data_v2,'Incident_Number':str(incident_number)})

@app.route('/exit', methods=['GET'])
def exit_chat():
    docx_filename = 'transcript.docx'
    document = Document()
    
    document.add_paragraph(f"Summary: {summary}")
    
    # Log user details
    for line in conversation:
        if line.startswith("Image:"):
            image_path = line.split("Image: ")[1].strip()
            if os.path.exists(image_path):
                try:
                    document.add_picture(image_path, width=docx.shared.Inches(6))
                except Exception as e:
                    print(f"Failed to load image {image_path}: {e}")  # Error handling
            else:
                print(f"Image path does not exist: {image_path}")  # Path check
        else:
            safe_line = line.encode('latin-1', 'replace').decode('latin-1')
            document.add_paragraph(safe_line)

    document.save(docx_filename)

    # Convert DOCX to HTML with images
    html_content = "<html><body>"
    for para in document.paragraphs:
        html_content += f"<p>{para.text}</p>"
        
        # Check for images in the paragraph
        for run in para.runs:
            if run.element.xpath('.//a:blip'):
                for blip in run.element.xpath('.//a:blip'):
                    image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if image_id:
                        image_path = document.part.related_parts[image_id].blob
                        # Save the image to a temporary location
                        image_filename = f'temp_image_{image_id}.png'  # Change the extension based on your image format
                        with open(image_filename, 'wb') as img_file:
                            img_file.write(image_path)
                        
                        # Add the image to HTML content
                        html_content += f'<img src="{image_filename}" style="max-width:100%; height:auto;"/><br>'

    html_content += "</body></html>"

    # Clear conversation
    conversation.clear()

    return jsonify({'content': html_content})

import re

import re
import os
from flask import request, send_file
from docx import Document

import re
import os
from flask import request, send_file
from docx import Document
import html  # Import the html module for decoding entities

@app.route('/download', methods=['POST'])
def download():
    edited_content = request.form['document_content']

    # Remove HTML tags from edited_content but retain <img> tags
    edited_content = re.sub(r'<br\s*\/?>', '\n', edited_content)  # Replace <br> with newlines
    edited_content = re.sub(r'(?i)<(?!img\s*).*?>', '', edited_content)  # Remove other HTML tags

    document = Document()

    # Split the content into lines and add each line as a paragraph
    for line in edited_content.splitlines():
        line = line.strip()
        if line:  # Add non-empty lines as paragraphs
            # Decode HTML entities
            line = html.unescape(line)

            # Check for <img> tags
            img_tags = re.findall(r'<img[^>]+src="([^"]+)"', line)
            for img_tag in img_tags:
                # If an image is found, add it to the document
                image_path = img_tag  # Adjust this if needed
                if os.path.exists(image_path):
                    document.add_picture(image_path, width=docx.shared.Inches(6))
            # Add the text without the <img> tag
            clean_line = re.sub(r'<img[^>]*>', '', line)  # Remove <img> tags from the line
            document.add_paragraph(clean_line)

    edited_docx_filename = 'edited_transcript.docx'
    document.save(edited_docx_filename)

    return send_file(edited_docx_filename, as_attachment=True)


import PyPDF2
import docx
from werkzeug.utils import secure_filename

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""  # Ensure empty text if extract_text returns None
            return text
    except Exception as e:
        return f"Error extracting text from PDF: {e}"

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"

# Function to extract text from TXT
def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, 'r') as file:
            return file.read()
    except Exception as e:
        return f"Error extracting text from TXT: {e}"

# Function to extract KYC-related information
def extract_kyc_info(text):
    try:
        name_match = re.search(r"Name:\s*(.*)", text)
        address_match = re.search(r"Address:\s*(.*)", text)
        dob_match = re.search(r"Date of Birth:\s*(\d{2}/\d{2}/\d{4})", text)

        kyc_info = {
            "name": name_match.group(1) if name_match else "N/A",
            "address": address_match.group(1) if address_match else "N/A",
            "dob": dob_match.group(1) if dob_match else "N/A"
        }
        return kyc_info
    except Exception as e:
        return {"error": f"Error extracting KYC info: {e}"}

@app.route('/analyze_document', methods=['POST'])
def analyze_document():
    if 'document' not in request.files:
        return jsonify({"success": False, "error": "No document uploaded"}), 400
    
    document = request.files['document']
    
    # Ensure the filename is safe
    document_path = os.path.join("uploads", secure_filename(document.filename))
    os.makedirs("uploads", exist_ok=True)  # Ensure upload directory exists
    document.save(document_path)
    
    # Extract text based on the file type
    file_extension = os.path.splitext(document.filename)[1].lower()
    text = ""

    try:
        if file_extension == ".pdf":
            text = extract_text_from_pdf(document_path)
        elif file_extension == ".docx":
            text = extract_text_from_docx(document_path)
        elif file_extension == ".txt":
            text = extract_text_from_txt(document_path)
        else:
            return jsonify({"success": False, "error": "Unsupported file type"}), 400
        
        if "Error" in text:
            return jsonify({"success": False, "error": text}), 500
    except Exception as e:
        return jsonify({"success": False, "error": f"Failed to process file: {e}"}), 500

    # Extract KYC-related information
    kyc_info = extract_kyc_info_v2(text)
    answer_html = kyc_info.replace('\n', '<br>')
    # Clean up the uploaded file
    os.remove(document_path)
    print(answer_html)
    # Return the extracted information
    return jsonify({
        "success": True,
        "kyc": answer_html
    })

@app.route('/chatbot_document', methods=['POST'])
def chatbot_document():
    data = request.get_json()

    if 'message' not in data:
        return jsonify({"error": "Missing 'message' in request data"}), 400

    user_message = data['message']
    # Integrate chatbot response logic here
    response = document_chat(user_message)
    answer_html = response.replace('\n', '<br>')
    return jsonify({"response": answer_html})

if __name__ == "__main__":
    app.run(debug=False)